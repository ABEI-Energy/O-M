'''

This script contains the functions to be called from streamlit.py

'''
import pandas as pd
import numpy as np
import locale as lc
import datetime as dt
import streamlit as st
import requests
import utm
from docx import Document
from PIL import Image
from docx.shared import Cm
import io

lc.setlocale(lc.LC_ALL,'es_ES.UTF-8')

month = dt.datetime.now().strftime("%B %Y").capitalize()
year = month.split(' ')[-1]


def excel_reader(directory, df_aux = None):

    if st.session_state.calculo_plantilla:

        xls = pd.ExcelFile(directory)

        hoja_portada = pd.read_excel(xls, 'Portada', index_col=None, header = None)
        hoja_tablas_plantilla = pd.read_excel(xls, 'Tablas', index_col=None, header = None)


        # Cogemos sólo los datos de portada que nos interesan

        df_portada = pd.DataFrame(hoja_portada[263:])
        df_portada.dropna(axis = 1, how = 'all', inplace = True)
        df_portada.dropna(axis = 0, how = 'all', inplace = True)
        df_portada.dropna(axis = 1, thresh = 10, inplace = True)

        df_portada = df_portada.iloc[:,0:-1]
        df_portada.dropna(axis = 0, how = 'any', inplace = True)
        df_portada.reset_index(inplace = True, drop = True)
        df_portada.columns = pd.RangeIndex(len(df_portada.columns))

        dict_portada = {}

        dict_portada['degTable'] = df_portada.iloc[3,1]*100
        dict_portada['mtProdTable'] = df_portada.iloc[4,1]
        dict_portada['btProdTable'] = df_portada.iloc[5,1]
        dict_portada['mtAggTable'] = df_portada.iloc[6,1]
        dict_portada['copRadTable'] = df_portada.iloc[7,1]
        dict_portada['copFilRadTable '] = df_portada.iloc[8,1]
        dict_portada['horRadTable'] = df_portada.iloc[9,1]
        dict_portada['monthPRTable'] = df_portada.iloc[10,1]
        dict_portada['monthAvailTable '] = df_portada.iloc[13,1] 
        dict_portada['dateTime'] = month

        # Get the tables.

        df_portada = hoja_tablas_plantilla.apply(lambda x: x.str.strip())
        end_tables = df_portada.loc[df_portada.iloc[:,0]=='TOTAL'].index #Empezar, la primera siempre en el mismo lado, y la segunda el primer index +2 
        df_portada = hoja_tablas_plantilla

        # We put in df_flagT1 the data for the first table, and we will build there.
        df_flagT1_aux = df_portada

        # Bloque 1
        df_flagT1_aux = df_portada

        df_flagT1_aux = df_flagT1_aux.iloc[:,:3]
        df_flagT1_aux = df_flagT1_aux.iloc[3:(end_tables[0]+1)]
        df_flagT1_aux.reset_index(inplace = True, drop = True)
        df_flagT1_aux.columns = ['Fecha', 'SET', 'Total Invs']
        df_flagT1_aux.name = 'PR.'
        
        st.session_state.flag_T1_aux = True

        # Bloque 2 va a flagT2
        df_flagT2 = df_portada
        cols_ct = df_flagT2.iloc[1].dropna()
        lenCT = len(df_flagT2.iloc[1].dropna())

        df_flagT2 = df_flagT2.iloc[:,3:3+lenCT]
        df_flagT2 = df_flagT2.iloc[3:(end_tables[0]+1)]
        df_flagT2.reset_index(inplace = True, drop = True)
        df_flagT2.columns = cols_ct
        df_flagT2.name = 'CTs.'

        df_flagT2 = pd.concat([df_flagT1_aux['Fecha'],df_flagT2], axis = 1)

        # flagT3 inversores
        df_flagT3 = df_portada
        cols_inv = df_flagT3.iloc[(end_tables[0]+2),:8]

        df_flagT3 = df_flagT3.iloc[:,:8] # Columnas
        df_flagT3 = df_flagT3.iloc[(end_tables[0]+6):(end_tables[1]+1)] # Filas
        df_flagT3.reset_index(inplace = True, drop = True)
        df_flagT3.columns = cols_inv
        df_flagT3.name = 'Irrad.'


        # flagT4 temperatura
        df_flagT4 = df_portada
        cols_temp = df_flagT4.iloc[(end_tables[0]+2),8:]

        df_flagT4 = df_flagT4.iloc[:,8:] # Columnas
        df_flagT4 = df_flagT4.iloc[(end_tables[0]+6):(end_tables[1]+1)] # Filas
        df_flagT4.reset_index(inplace = True, drop = True)
        df_flagT4.columns = cols_temp
        df_flagT4.name = 'Temp.'

        df_flagT4 = pd.concat([df_flagT1_aux['Fecha'],df_flagT4], axis = 1)


        st.session_state.calculo_plantilla = False
        st.session_state.plantilla = True

        return df_flagT1_aux, df_flagT2, df_flagT3, df_flagT4, dict_portada
        
    elif st.session_state.calculo_disponibilidad:

        df_flagT1_aux = df_aux

        xls = pd.ExcelFile(directory)

        hoja_calculo_disp = pd.read_excel(xls, 'Cálculo Disp. 24 (corr)', index_col=None, header = None)

        # Get the month we are
        st.session_state.month = hoja_calculo_disp.values[7][0].strftime('%B')

        # Ver cuántos días tiene el mes, esa columna varía en longitud.
        end_tables = hoja_calculo_disp.loc[hoja_calculo_disp.iloc[:,42]=='TOTAL'].index[0] #Empezar, la primera siempre en el mismo lado, y la segunda el primer index +2 
        df_calculo_disp = hoja_calculo_disp.iloc[52:end_tables + 1,42:44].reset_index(drop = True)
        df_calculo_disp.columns = ['Fecha', 'Availability']

        df_flagT1_aux = df_flagT1_aux.merge(df_calculo_disp,on = 'Fecha')
        df_flagT1_aux.iloc[-1,-1] = hoja_calculo_disp.iloc[1,14]

        df_flagT1_aux['Availability'] = df_flagT1_aux['Availability']*100
        # df_flagT1_aux['Availability'] = df_flagT1_aux['Availability'].astype(str) + ' %'

        st.session_state.days = end_tables - 52
        st.session_state.calculo_disponibilidad = False
        st.session_state.flag_T1_aux1 = False

        st.session_state.flag_T1_aux2 = True
        st.session_state.disponibilidad = True


        return df_flagT1_aux        

    elif st.session_state.calculo_pr:

        df_flag_T1_aux2 = df_aux

        xls = pd.ExcelFile(directory)

        hoja_calculo_disp = pd.read_excel(xls, 'Calculos Prod.  (corr)', index_col=None, header = None)

        # Ya sabemos cuántos días hay, porque aquí la hoja se va a poner llenando sin más.
        end_tables = st.session_state.days #Empezar, la primera siempre en el mismo lado, y la segunda el primer index +2 
        df_calculo_pr = hoja_calculo_disp.iloc[17:18 + st.session_state.days, 20:22].reset_index(drop = True)
        df_calculo_pr.columns = ['Fecha', 'PR']
        df_calculo_pr.iloc[-1,-1] = hoja_calculo_disp.iloc[19,25]
        df_calculo_pr.iloc[-1,0] = 'TOTAL'

        df_flag_T1_aux2 = df_flag_T1_aux2.merge(df_calculo_pr,on = 'Fecha')

        df_flag_T1_aux2['PR'] = df_flag_T1_aux2['PR']*100
        # df_flag_T1_aux2['PR Nuevo'] = df_flag_T1_aux2['PR Nuevo'].astype(str) + ' %'

        st.session_state.calculo_disponibilidad = False
        st.session_state.df_flag_T1_aux2 = False

        st.session_state.tablesDone = True        

        st.session_state.calculo_pr = False
        st.session_state.flag_T1_aux2 = False

        st.session_state.pr = True


        return df_flag_T1_aux2        

def duplicateDoc():

    filemodelo = 'Resources/model/Modelo informe cliente.docx'

    return Document(filemodelo)


def insert_image_in_cell(doc, picDict):

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if ("flagProductionSETFig" in cell.text):
                    imagen = Image.open(picDict['SetProduction'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))
                if ("flagProductionCTSFig" in cell.text):
                    imagen = Image.open(picDict['CTSProduction'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))
                if ("flagProductionRadiationFig" in cell.text):
                    imagen = Image.open(picDict['CopHorizRadiation'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))  
                if ("flagTemperaturesFig" in cell.text):
                    imagen = Image.open(picDict['Temperatures'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))  
                if ("flagPRMonthFig" in cell.text):
                    imagen = Image.open(picDict['PR'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))      

                if ("flagAvailabilityFig" in cell.text):
                    imagen = Image.open(picDict['Availability'])
                    size = imagen.size
                    imWidth = float(size[0])
                    imHeight = float(size[1])                    
                    if imWidth > 17.00:
                        imHeight = imHeight*17.00/imWidth
                        imWidth = 17.00
                    image_io = io.BytesIO()
                    imagen.save(image_io, format = 'PNG')
                    image_io.seek(0)
                    cell.text = ""
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    run.add_picture(image_io, width = Cm(imWidth), height = Cm(imHeight))                                                     
    
    st.session_state.picsDone = True

def docWriter(docxFile,docxDict):

    #Headers
    for section in docxFile.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for word in docxDict:
                            if word in paragraph.text:
                                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                                paragraph.style = docxFile.styles['headerStyle1'] 

    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    previousStyle = docxFile.styles['tablePortada']
            
                    for word in docxDict:
                        if word in paragraph.text:
                            previousStyle = paragraph.style.name
                            paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                            paragraph.style = docxFile.styles[previousStyle]   
                            table.style = docxFile.styles['tablePortadas']
    
    #Resto del documento
    for paragraph in docxFile.paragraphs:
        for word in docxDict:
            if word in paragraph.text:
                previousStyle = paragraph.style.name
                paragraph.text = paragraph.text.replace(word,str(docxDict[word]))
                paragraph.style = docxFile.styles[previousStyle] 

    st.session_state.wordsDone = True
    st.session_state.tablesDone = True

def docTabler(docxFile, df_flagT1, df_flagT2, df_flagT3, df_flagT4):

    for table in docxFile.tables:
        for row in table.rows:
            for cell in row.cells:
                if "flagT1" in cell.text:
                    cell.text = "Fecha"
                    cell.paragraphs[0].style  = docxFile.styles['tableVal']
                    for i in range(len(df_flagT1)):
                        table.add_row()

                    for i in range(df_flagT1.shape[0]):
                        for j in range(df_flagT1.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df_flagT1.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = docxFile.styles['tableVal']

                    for col in table.columns:
                        col.width = Cm(2)

                    continue

                if "flagT2" in cell.text:
                    cell.text = "Fecha"
                    cell.paragraphs[0].style = docxFile.styles['tableVal']
                    for i in range(len(df_flagT2)):
                        table.add_row()

                    for i in range(df_flagT2.shape[0]):
                        for j in range(df_flagT2.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df_flagT2.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = docxFile.styles['tableVal']

                    for col in table.columns:
                        col.width = Cm(1.65)

                    continue


                if "flagT3" in cell.text:
                    cell.text = "Fecha"
                    cell.paragraphs[0].style = docxFile.styles['tableVal']
                    for i in range(len(df_flagT3)):
                        table.add_row()

                    for i in range(df_flagT3.shape[0]):
                        for j in range(df_flagT3.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df_flagT3.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = docxFile.styles['tableVal']
            
                    for col in table.columns:
                        col.width = Cm(1.7)

                    continue

                if "flagT4" in cell.text:
                    cell.text = "Fecha"
                    cell.paragraphs[0].style = docxFile.styles['tableVal']
                    for i in range(len(df_flagT4)):
                        table.add_row()

                    for i in range(df_flagT4.shape[0]):
                        for j in range(df_flagT4.shape[-1]):
                            table.cell(i+1,j).paragraphs[0].text = str(df_flagT4.values[i,j])
                            table.cell(i+1,j).paragraphs[0].style = docxFile.styles['tableVal']

                    for col in table.columns:
                        col.width = Cm(1.9)

                    continue


    st.session_state.tablerDone = True